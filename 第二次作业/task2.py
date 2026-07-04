import pandas as pd
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==========================================
# 准备工作与数据加载
# ==========================================
# 考虑到文件在第一次作业的文件夹中，使用相对路径获取数据
data_path = "第一次作业/600519_history_data.csv"

if not os.path.exists(data_path):
    print(f"找不到数据文件！请确认你的代码是在'第二次作业'文件夹中运行，并且数据文件位于 {data_path}")
    exit()

df = pd.read_csv(data_path)
# 确保日期按升序排列以便计算指标
df['date'] = pd.to_datetime(df['date'])
df = df.sort_values('date').reset_index(drop=True)

plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False

# ==========================================
# 1. 基础诊断分析
# ==========================================
missing_values = df.isnull().sum()
desc_stats = df.describe().round(2)

# ==========================================
# 2. 指标计算
# ==========================================
# (1) RSI (14日)
delta = df['close'].diff()
up = delta.clip(lower=0)
down = -1 * delta.clip(upper=0)
ema_up = up.ewm(com=13, adjust=False).mean()
ema_down = down.ewm(com=13, adjust=False).mean()
rs = ema_up / ema_down
df['RSI_14'] = 100 - (100 / (1 + rs))

# (2) MACD
exp1 = df['close'].ewm(span=12, adjust=False).mean()
exp2 = df['close'].ewm(span=26, adjust=False).mean()
df['DIF'] = exp1 - exp2
df['DEA'] = df['DIF'].ewm(span=9, adjust=False).mean()
df['MACD'] = 2 * (df['DIF'] - df['DEA'])

# (3) Bollinger Bands (20日)
df['BOLL_MID'] = df['close'].rolling(window=20).mean()
df['BOLL_STD'] = df['close'].rolling(window=20).std()
df['BOLL_UPPER'] = df['BOLL_MID'] + 2 * df['BOLL_STD']
df['BOLL_LOWER'] = df['BOLL_MID'] - 2 * df['BOLL_STD']

# (4) 扩展指标 ATR (14日)
df['prev_close'] = df['close'].shift(1)
tr1 = df['high'] - df['low']
tr2 = (df['high'] - df['prev_close']).abs()
tr3 = (df['low'] - df['prev_close']).abs()
df['TR'] = pd.concat([tr1, tr2, tr3], axis=1).max(axis=1)
df['ATR_14'] = df['TR'].rolling(window=14).mean()

# ==========================================
# 3. 绘制并保存可视化图形
# ==========================================
# 图1：RSI
plt.figure(figsize=(10, 4))
plt.plot(df['date'], df['RSI_14'], label='RSI (14)', color='blue')
plt.axhline(70, color='red', linestyle='--', alpha=0.5)
plt.axhline(30, color='green', linestyle='--', alpha=0.5)
plt.title('图1：贵州茅台 RSI 走势图')
plt.legend()
plt.tight_layout()
plt.savefig('rsi_chart.png')
plt.close()

# 图2：MACD
plt.figure(figsize=(10, 4))
plt.plot(df['date'], df['DIF'], label='DIF', color='blue')
plt.plot(df['date'], df['DEA'], label='DEA', color='orange')
plt.bar(df['date'], df['MACD'], label='MACD Histogram', color=np.where(df['MACD']>0, 'red', 'green'), alpha=0.5)
plt.title('图2：贵州茅台 MACD 走势图')
plt.legend()
plt.tight_layout()
plt.savefig('macd_chart.png')
plt.close()

# 图3：布林带
plt.figure(figsize=(10, 5))
plt.plot(df['date'], df['close'], label='收盘价', color='black')
plt.plot(df['date'], df['BOLL_UPPER'], label='上轨 (Upper)', color='red', linestyle='--')
plt.plot(df['date'], df['BOLL_MID'], label='中轨 (Middle)', color='blue', linestyle='--')
plt.plot(df['date'], df['BOLL_LOWER'], label='下轨 (Lower)', color='green', linestyle='--')
plt.fill_between(df['date'], df['BOLL_UPPER'], df['BOLL_LOWER'], color='gray', alpha=0.1)
plt.title('图3：贵州茅台布林带 (Bollinger Bands) 走势图')
plt.legend()
plt.tight_layout()
plt.savefig('boll_chart.png')
plt.close()

# 图4：ATR
plt.figure(figsize=(10, 4))
plt.plot(df['date'], df['ATR_14'], label='ATR (14)', color='purple')
plt.title('图4：贵州茅台 ATR 真实波动幅度走势图')
plt.legend()
plt.tight_layout()
plt.savefig('atr_chart.png')
plt.close()

# ==========================================
# 4. 生成符合格式要求的 Word 文档
# ==========================================
print("正在生成排版完美的 Word 文档，请稍候...")
doc = Document()

def add_formatted_paragraph(document, text, is_bold=False):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.add_run(text)
    run.bold = is_bold
    run.font.name = 'SimSun'
    run.font.size = Pt(10.5) # 五号字
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p

# 写入第一题内容
add_formatted_paragraph(doc, "1. 数据诊断分析结果", is_bold=True)
add_formatted_paragraph(doc, "（1）缺失值检查：")
missing_text = ", ".join([f"{idx}: {val}" for idx, val in missing_values.items()])
add_formatted_paragraph(doc, f"经检验，数据缺失情况如下：{missing_text}。数据完整度高，无需剔除或插值修补。")
add_formatted_paragraph(doc, "（2）收盘价（close）描述性统计量：")
close_stats = desc_stats['close']
add_formatted_paragraph(doc, f"样本量: {close_stats['count']}，均值: {close_stats['mean']}，标准差: {close_stats['std']}，"
                             f"最小值: {close_stats['min']}，最大值: {close_stats['max']}。")

add_formatted_paragraph(doc, "") # 空行

# 写入第二题内容
add_formatted_paragraph(doc, "2. 基础描述指标分析", is_bold=True)
add_formatted_paragraph(doc, "（1）RSI (相对强弱指标)：计算N日内上涨幅度之和的平均值与下跌幅度绝对值之和的平均值的比率。作用是衡量资产价格变动的速度和幅度，常以70以上作为超买信号，30以下作为超卖信号。")
add_formatted_paragraph(doc, "（2）MACD (平滑异同移动平均线)：由12日与26日EMA差值求出DIF，再求DIF的9日EMA得出DEA。作用是判断趋势。DIF上穿DEA视为金叉买入，下穿视为死叉卖出。")
add_formatted_paragraph(doc, "（3）Bollinger Bands (布林带)：以20日均线为中轨，上下各加减两倍标准差构成通道。作用是衡量市场波动率，价格触碰上轨暗示超买，触碰下轨暗示超卖。")

add_formatted_paragraph(doc, "") 

# 写入扩展指标（第四题）
add_formatted_paragraph(doc, "4. 扩展指标介绍：ATR (真实波动幅度)", is_bold=True)
add_formatted_paragraph(doc, "计算方法：先计算每日的真实波幅（TR），即最高价减最低价、最高价减前收盘价绝对值、最低价减前收盘价绝对值三者中的最大值。再取其14日移动平均。")
add_formatted_paragraph(doc, "作用：ATR 不预测价格方向，专门用于反映市场价格波动的剧烈程度。在量化交易中主要用于风控、动态设置止损位以及仓位管理。")

add_formatted_paragraph(doc, "") 

# 写入图表与解读
add_formatted_paragraph(doc, "3. 可视化图表及解读", is_bold=True)

# 插入 RSI
doc.add_picture('rsi_chart.png', width=Inches(6.0))
add_formatted_paragraph(doc, "图1解读：RSI 曲线反映了茅台股价在近一年内的超买超卖情况。当红线（70）被突破时，说明短期动能过热；触及绿线（30）时，说明存在反弹可能。")

# 插入 MACD
doc.add_picture('macd_chart.png', width=Inches(6.0))
add_formatted_paragraph(doc, "图2解读：MACD 图中，红绿柱状图的高低展示了多空趋势的强度。蓝线(DIF)和橙线(DEA)的交叉点清晰地指示了潜在的波段买卖点。")

# 插入 BOLL
doc.add_picture('boll_chart.png', width=Inches(6.0))
add_formatted_paragraph(doc, "图3解读：布林带图直观展现了股价的运行通道。多数时间收盘价在上下轨之间震荡；通道变窄时往往意味着变盘在即。")

# 插入 ATR
doc.add_picture('atr_chart.png', width=Inches(6.0))
add_formatted_paragraph(doc, "图4解读：ATR 曲线捕捉了标的波动率的变化。峰值阶段说明当时市场情绪激烈、日内振幅大；低谷阶段说明交投相对清淡。")

# 强制全部段落首行缩进及字体确认
for para in doc.paragraphs:
    para.paragraph_format.first_line_indent = Pt(21)

doc.save('姓名_TASK2.docx')
print("执行完毕！文档已保存为 '姓名_TASK2.docx'")